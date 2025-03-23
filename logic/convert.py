import os
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def parse_table_lines(lines):
    table = []
    for line in lines:
        if '|' in line:
            row = [cell.strip() for cell in line.strip().split('|')[1:-1]]
            table.append(row)
    return table

def convert_zip_to_word_only(input_folder, selected_files, title, author, date, work_dir):
    doc = Document()

    # Couverture
    doc.add_heading(title, 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"{author} – {date}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph().add_run(" ").add_break()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Segoe UI'
    font.size = Pt(11)

    media_path = os.path.join(input_folder, 'media')

    for md_file in selected_files:
        page_name = os.path.splitext(md_file)[0]
        md_path = os.path.join(input_folder, md_file)

        # Détection niveau : si le nom contient "/", on le traite comme sous-niveau
        sections = page_name.split('/')
        for i, section in enumerate(sections):
            doc.add_heading(section.replace('-', ' '), level=i+1)

        with open(md_path, "r", encoding="utf-8") as f:
            lines = f.readlines()

        buffer = []
        in_table = False

        for line in lines:
            line = line.strip()

            # Images ![alt](media/image.png)
            if line.startswith("![](") and ")" in line:
                image_file = line.split("![](")[-1].split(")")[0]
                image_path = os.path.join(input_folder, image_file)
                if os.path.exists(image_path):
                    try:
                        doc.add_picture(image_path, width=Inches(5.5))
                        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception:
                        doc.add_paragraph(f"[Image introuvable ou non supportée : {image_file}]")
                continue

            # Début de table
            if '|' in line and line.count('|') >= 2:
                in_table = True
                buffer.append(line)
                continue

            if in_table:
                if line == '' or not ('|' in line):
                    # fin de la table
                    table = parse_table_lines(buffer)
                    if table:
                        rows, cols = len(table), len(table[0])
                        tbl = doc.add_table(rows=rows, cols=cols)
                        tbl.style = 'Table Grid'
                        for i, row in enumerate(table):
                            for j, cell in enumerate(row):
                                tbl.cell(i, j).text = cell
                    buffer = []
                    in_table = False
                else:
                    buffer.append(line)
                continue

            # Titres markdown
            if line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("- "):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif re.match(r'^\d+\. ', line):
                doc.add_paragraph(line, style='List Number')
            elif line == "":
                doc.add_paragraph("")
            else:
                doc.add_paragraph(line)

        doc.add_page_break()

    docx_output = os.path.join(work_dir, "Notion_Document.docx")
    doc.save(docx_output)
    return docx_output